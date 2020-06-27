#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Analysis of COVID19 spread from public data
released by covid19.org for Indian data and 
John Hopkins University data for global data

@author: Abhijit Bhattacharyya
         Nuclear Physics Division
         Bhabha Atomic Research Centre
         Mumbai 400 085
         EMAIL:abhihere@gmail.com, vega@barc.gov.in
@File: covid19VEGA02.py
@author: Abhijit Bhattacharyya
Created on Tue May 26 23:29:22 2020
"""

import sys
import os 
import gc
from tabulate import tabulate
import argparse
import pandas as pd
from pandas.plotting import table
import numpy as np
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from mpl_toolkits.mplot3d import Axes3D
from scipy.optimize import curve_fit
from matplotlib import ticker
import pycountry_convert as pc
import folium
import branca
import datetime as dt
from datetime import datetime, timedelta, date
from scipy.interpolate import make_interp_spline, BSpline
import weasyprint as wsp
import PIL as pil
import docx
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

import plotly.graph_objects as go
import plotly.express as px
import plotly.io as pio
from plotly.subplots import make_subplots
import cufflinks as cf
from plotly.offline import download_plotlyjs, init_notebook_mode, plot, iplot
pio.templates.default = "xgridoff"

import glob, json, requests
import calmap
from bs4 import BeautifulSoup

from keras.layers import Input, Dense, Activation, LeakyReLU
from keras import models
from keras.optimizers import RMSprop, Adam

from statsmodels.tsa.statespace.sarimax import SARIMAX
from statsmodels.tsa.arima_model import ARIMA
from random import random
from tqdm import tqdm

from sklearn import preprocessing
from sklearn.model_selection import KFold

import warnings
warnings.filterwarnings('ignore')

# -----------------------------------------------------------------------
# -------------------  FUNCTION definitions ----------------------------
out = ""#+"output/"
cmdParser = argparse.ArgumentParser() 

def RMSLE(pred, actual):
    return np.sqrt(np.mean(np.power((np.log(pred + 1) - np.log(actual + 1)), 2)))

def _convert_date_str(df):
    try:
        df.columns = list(df.columns[:4]) + [datetime.strptime(d, "%m/%d/%y").date().strftime("%Y-%m-%d") for d in df.columns[4:]]
    except:
        print('_convert_date_str failed with %y, try %Y')
        df.columns = list(df.columns[:4]) + [datetime.strptime(d, "%m/%d/%Y").date().strftime("%Y-%m-%d") for d in df.columns[4:]]


# ----------- ------------------------- Write a Document
def makeDoc(thisDF, docHead, fName):
    # myDoc = docx.Document(docFName)   # to read old file
    myDoc = docx.Document()             # to create new file
    myDoc.core_properties.author = 'Abhijit Bhattacharyya'
    myDoc.core_properties.owner = 'Abhijit Bhattacharyya'
    myDoc.core_properties.title = docHead
    thisDF.reset_index(drop=False, inplace=True)
    #myDoc.add_heading('  ',2)
    
    currSection = myDoc.sections[-1]
   # currSection.orientation = WD_ORIENT.LANDSCAPE
    currSection.orientation = WD_ORIENT.PORTRAIT
    
    myDocHead = myDoc.add_heading(docHead,0)
    myDocHead.bold = True
    #myDocHead.underline = True
    tt = myDoc.add_table(thisDF.shape[0] + 1, thisDF.shape[1])   # row, col
    tt.style = 'LightShading-Accent1'
    tt.autofit = False

    for j in range(thisDF.shape[-1]):
        tt.cell(0, j).text = thisDF.columns[j]      # header
    
# cell width in EMU i.e. 1 inch = 914400 EMU.
        
    for i in range(thisDF.shape[0]):        
        for j in range(thisDF.shape[-1]):                                            
            tt.cell(i+1, j).text = str(thisDF.values[i, j])  # rest part of the dataframe
        tt.cell(i+1, 0).width =  2560320            
        tt.cell(i+1, 1).width =  4105656            
        tt.cell(i+1, 2).width =  5678424            
        tt.cell(i+1, 3).width =  7397496            
        tt.cell(i+1, 4).width =  9509760            
        tt.cell(i+1, 5).width =  11420856
        
    myDoc.save(fName)
    # Working


#------------ -------------------------  Plot Global Figures 
def plotGlobalGig(thisDF, threshold = None, gType = None, getCNF = False, nCNFIndia = None):
    f = plt.figure(figsize = (20, 12))
    ax = f.add_subplot(111)
    
    dateList = thisDF.columns.tolist()
    #xList = [(dt.datetime.strptime(d, '%m/%d/%y').date()) for d in dateList]
    xList = dateList
    colorList = ['crimson', 'orange', 'tomato', 'moccasin', 'darkviolet', 'lightseagreen', 
                 'plum','steelblue', 'navy', 'palegreen', 'forestgreen', 'limegreen'
                 ]
    jColor = 0
    for ii, Country in enumerate(thisDF.index):        
        if (ii > 9):
            if (Country != 'India' and Country != 'China'):
                continue
        
        gT1 = thisDF.loc[thisDF.index == Country].values[0]
        if (getCNF == True):
            if Country != "India":
                continue
            else:
                return gT1[-1]
        date0 = ((thisDF.loc[thisDF.index == Country]) > threshold).idxmax(axis=1)[0]
        # date0 = dt.datetime.strptime(date0, '%m/%d/%y').date()        
        posDate0 = pd.Index(xList).get_loc(date0)
                                
        xnew = xList[posDate0:]
        gT1 = gT1[gT1 > threshold]
        yVals = gT1
        
        if (getCNF == False):
            percentage = f"  ( {((gT1[-1] / nCNFIndia) ): .2%})"
        if (gType == "Confirmed"):
            testCountry = str(ii+1) + "  " + Country + ": " + str(gT1[-1])
        else:
            if (Country == 'India'):
                testCountry = str(ii+1) + " " + Country + ": " + str(gT1[-1]) + percentage
            else:
                testCountry = str(ii +1) + "  " + Country + ": " + str(gT1[-1])
        

        strTitle = "Trend comparison of some countries with India for " + gType + " cases with threshold " + str(threshold)
        strYlab = "Number of " + gType + " cases"            
                
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%m/%d/%Y'))
        plt.gca().xaxis.set_major_locator(mdates.DayLocator(interval = 7))
        plt.gcf().autofmt_xdate()
                
        if (Country != 'India'):
            plt.plot(xnew, yVals, '-o', color = colorList[jColor], label = testCountry, linewidth = 3, markevery = [-1])
        else:
            marker_style = dict(linewidth = 3, linestyle = '-', marker = 'o', markersize = 8, markerfacecolor = colorList[jColor]) #markerfacecolor = '#ffffff')
            plt.plot(xnew, yVals, "-.", label = testCountry, **marker_style)
        jColor += 1    
        

                
        plt.xticks(rotation=70)
        plt.yscale("log")
        plt.xlabel("Date", fontsize = 17)
        plt.ylabel(strYlab, fontsize = 17)
        plt.title(strTitle, fontsize = 15)        
        plt.legend(loc = 0)
        plt.grid(which="both")        
        
        # reference lines
        # ------------------------   reference for case doubling every day
        print(xList) #, "   ", xList[0]+timedelta(5))
        exit(0)
        
        x = [xList[0] + timedelta(days = xValue) for xValue in range (21)]
        xd = [(xi - x[0]).days for xi in x]
        y = [2**(xdi + np.log2(threshold)) for xdi in xd]
        plt.plot(x, y, "--", linewidth = 2, color="gray")
        plt.annotate("Doubling every day", (x[-21], y[-1]), xycoords = "data", fontsize = 12, alpha = 0.5)
        
        # ------------------------   reference for case doubling every 4th day
        x = [xList[0] + timedelta(days = xValue) for xValue in range (80)]
        xd = [(xi - x[0]).days for xi in x]
        y = [2**(xdi / 4 + np.log2(threshold)) for xdi in xd]
        plt.plot(x, y, "--", linewidth = 2, color="gray")
        plt.annotate("Doubling every 4th day", (x[-25], y[-1]), xycoords = "data", fontsize = 12, alpha = 0.5)
   
        # ------------------------   reference for case doubling every week
        x = [xList[0] + timedelta(days = xValue) for xValue in range (123)]
        xd = [(xi - x[0]).days for xi in x]
        y = [2**(xdi / 7 + np.log2(threshold)) for xdi in xd]
        plt.plot(x, y, "--", linewidth = 2, color="red")
        plt.annotate("Doubling every week", (x[-20], y[-1]), color="red", xycoords = "data", fontsize = 12, alpha = 0.5)
   
        # ------------------------   reference for case doubling every fortnight
        x = [xList[0] + timedelta(days = xValue) for xValue in range (130)]
        xd = [(xi - x[0]).days for xi in x]
        y = [2**(xdi / 14 + np.log2(threshold)) for xdi in xd]
        plt.plot(x, y, "--", linewidth = 2, color="red")
        plt.annotate("Doubling every fortnight", (x[-20], y[-1]), color="red", xycoords = "data", fontsize = 12, alpha = 0.5)
        
        
        
#------------------------------- Get World Data over Net starts 
def getDataOverNet():
    print("\n You have chosen to download the latest Data from the internet....\n")
    req = requests.get('https://www.worldometers.info/coronavirus/')
    soup = BeautifulSoup(req.text, "lxml")
    
    df_country = soup.find('div',attrs={"id" : "nav-tabContent"}).find(
        'table',attrs={"id" : "main_table_countries_today"}).find_all('tr')
    arrCountry = []

    for i in range(8,len(df_country)-1):
        tmp = df_country[i].find_all('td')
        country=""
    
        if (tmp[0].string) is not  None:        
            if (tmp[0].string.find('<a') == -1):
                country = [tmp[0].string]
            else:
                country = [tmp[0].a.string] # Country
        else:
            continue
    

        for j in range(1,13):
            if (str(tmp[j].string) == 'None' or str(tmp[j].string) == ' '):
                country = country + [0]
            else:            
                country = country + [(tmp[j].string.replace(',','').replace('+',''))]   
    
        #print(country)   # diagnosis of retrieved lines population could not be found as it was relinked again and so ommitted
        country.pop(0)
        arrCountry.append(country)
        
        df_worldinfor = pd.DataFrame(arrCountry)
        df_worldinfor.columns = ['Country','Total Cases','New Cases','Total Deaths','New Deaths','Total Recovered','Active Cases',
                         'Serious Critical','Total Cases/1M pop','Deaths/1M pop','Total Test','Tests/1M pop']
                         #'Population','Continent']

    for i in range(0,len(df_worldinfor)):
        df_worldinfor['Country'].iloc[i] = df_worldinfor['Country'].iloc[i].strip()

    df_worldinfor = df_worldinfor.set_index('Country')

    # Saving the donwloaded data in CSV file for future use
    df_worldinfor.to_csv('DATASET/worldometers_info.csv')
    return df_worldinfor

#------------------------------- Get World Data over Net ends --------------------------------

# ---------------------------------   FUNCTION DEF ENDS HERE ---------------------------------


# -------------------------  Process comamnd line arguments ------------------
opt1 = opt2 = opt3 = 0
if (((len(sys.argv) > 1) and (sys.argv[1] == '-h')) or (len(sys.argv) == 1)):
    print("Command: python covid19Vega02.py -o <opt1>")
    print(" opt1: 0  - read stored data and continue processing....")
    print(" opt1: 1  - download data from net, store to disk and continue processing.")
    print(" opt1: 2  - download data from net and EXIT without processing.") 
    print("\n")

    if (len(sys.argv) == 1):
        print(" ........ please run again with options.")
    sys.exit(0)
    
if (len(sys.argv) > 1):
    cmdParser.add_argument('-o', type = str)
    cmdArguments = cmdParser.parse_args()
    optList = cmdArguments.o.split(',')

if (int(optList[0]) >= 2):
    getDataOverNet()
    exit(0)
elif (int(optList[0]) == 1):   # means opt1 = 1 i.e. get data from net and store
    df_worldinfor = getDataOverNet()
elif (int(optList[0]) <= 0):
    root_path_world_data = './DATASET'
    worldometerData = f'{root_path_world_data}/worldometers_info.csv'
    # df_worldinfor = pd.read_csv(worldometerData, index_col='Country') #None)    
    #df_worldinfor = pd.read_csv(worldometerData, index_col='None')    
    df_worldinfor = pd.read_csv(worldometerData)    
    for i in range(0,len(df_worldinfor)):
        df_worldinfor['Country'].iloc[i] = df_worldinfor['Country'].iloc[i].strip()
    
# ------------------------ Command line processing ends here---------------------------------------


# --------------------------------------- Main code starts here ---------------------------------------------
df_worldinfor1=df_worldinfor.apply(pd.to_numeric,errors='ignore')
df_worldinfor1['Total Recovered']= pd.to_numeric(df_worldinfor1['Total Recovered'],errors='coerce')
df_worldinfor1 = df_worldinfor1[df_worldinfor1.Country != 'Total:']
df_worldinfor1.index = df_worldinfor1['Country']
df_worldinfor1 = df_worldinfor1.drop(['Country'], axis=1)

# Validating testing data around the world **************************************

df_testing1 = df_worldinfor1.drop(['Total Cases','New Cases','Total Deaths','New Deaths','Total Recovered','Active Cases',
                                'Serious Critical','Total Cases/1M pop','Deaths/1M pop'], axis = 1)

df_testingS1 = df_testing1.sort_values('Tests/1M pop', ascending=False)
#print(tabulate(df_testingS1, headers = df_testingS1.columns.tolist(), tablefmt='psql'))

fig = plt.figure(figsize=(15, 11))
fig.add_subplot(111)
#plt.axes(axisbelow=True)
#plt.barh(df_testingS1.sort_values('Tests/1M pop')["Tests/1M pop"].index[-50:], 
#         df_testingS1.sort_values('Tests/1M pop')["Tests/1M pop"].values[-50:], color="crimson")
plt.bar(df_testingS1.sort_values('Tests/1M pop')["Tests/1M pop"].index[-50:], 
         df_testingS1.sort_values('Tests/1M pop')["Tests/1M pop"].values[-50:], color="crimson")
plt.tick_params(size=5, labelsize=13)
plt.xticks(rotation=70)
plt.ylabel("Test/1M pop", fontsize=18)
plt.title("Top Countries (Tests / 1M pop )", fontsize=20)
plt.grid(alpha = 0.2)
plt.tight_layout()
plt.savefig(out+'PLOT/code2/testsPer1M.png')
# plt.show()
plt.close()

# Computing mortality Rate on Testing *******************************************

df_testing2 = df_worldinfor1.drop(
    ['New Cases','New Deaths','Total Recovered','Active Cases','Serious Critical','Deaths/1M pop','Total Cases/1M pop'], axis = 1)

df_testing2["MortalityRate"] = np.round(100 * df_testing2["Total Deaths"] / df_testing2["Total Cases"], 2)
df_testing2["Positive"] = np.round(100 * df_testing2["Total Cases"] / df_testing2["Total Test"], 2)

#df_testing2 = df_testing2.sort_values('Total Cases', ascending=False)
df_testing2 = df_testing2.sort_values('Country', ascending=True)

df_testing2.style.background_gradient(cmap='Blues',subset=["Total Test"])\
                        .background_gradient(cmap='Reds',subset=["Tests/1M pop"])\
                        .background_gradient(cmap='Greens',subset=["Total Cases"])\
                        .background_gradient(cmap='Purples',subset=["Total Deaths"])\
                        .background_gradient(cmap='YlOrBr',subset=["MortalityRate"])\
                        .background_gradient(cmap='bone_r',subset=["Positive"])

#print(tabulate(df_testing2, headers = df_testing2.columns.tolist(), tablefmt='psql'))

myDocHead = 'Country-wise tests and mortality report'
docFName = 'PLOT/code2/mortalityTable.docx'
#  makeDoc(df_testing2, myDocHead, docFName)    ## WORKING

#------------------------------------------------------------------------------------------
#----------------------- JHU DATA --------------------------------------------
JHUconfirmed = pd.read_csv(
    'DATASET/JohnHopkinsU_CSSE/csse_covid_19_data/csse_covid_19_time_series/time_series_covid19_confirmed_global.csv',
    parse_dates=True, index_col=None
    )

JHUdeaths = pd.read_csv(
    'DATASET/JohnHopkinsU_CSSE/csse_covid_19_data/csse_covid_19_time_series/time_series_covid19_deaths_global.csv',
    parse_dates=True, index_col=None
    )

JHUrecovered = pd.read_csv(
    'DATASET/JohnHopkinsU_CSSE/csse_covid_19_data/csse_covid_19_time_series/time_series_covid19_recovered_global.csv',
    parse_dates=True, index_col=None
    )


_convert_date_str(JHUconfirmed)
_convert_date_str(JHUdeaths)
_convert_date_str(JHUrecovered)


JHUconfirmed.replace(np.nan,'', regex=True)
JHUdeaths.replace(np.nan,'', regex=True)
JHUrecovered.replace(np.nan,'', regex=True)

JHUconfirmed = JHUconfirmed.rename(columns={"Province/State":"State", "Country/Region":"Country"}) #, inplace=True)
JHUdeaths = JHUdeaths.rename(columns={"Province/State":"State", "Country/Region":"Country"})       #, inplace=True)
JHUrecovered = JHUrecovered.rename(columns={"Province/State":"State", "Country/Region":"Country"}) #, inplace=True) 

#JHUconfirmed = JHUconfirmed[~JHUconfirmed]


JHUconfirmed1=pd.melt(JHUconfirmed,id_vars=['State','Country','Lat','Long'],var_name='date', value_name='Confirmed')
JHUconfirmed1['date'] = JHUconfirmed1['date'].astype('datetime64[ns]') 
JHUconfirmed1.sort_values(by="date")

JHUdeaths1=pd.melt(JHUdeaths,id_vars=['State','Country','Lat','Long'],var_name='date', value_name='Deceased')
JHUdeaths1['date'] = JHUdeaths1['date'].astype('datetime64[ns]') 
JHUdeaths1.sort_values(by="date")

JHUrecovered1=pd.melt(JHUrecovered,id_vars=['State','Country','Lat','Long'],var_name='date', value_name='Recovered')
JHUrecovered1['date'] = JHUrecovered1['date'].astype('datetime64[ns]') 
JHUrecovered1.sort_values(by="date")


#JHUActive = JHUconfirmed - JHUdeaths - JHUrecovered
#JHUActive1 = JHUconfirmed1 - JHUdeaths1 - JHUrecovered1
# print(JHUconfirmed.shape,'   ', JHUdeaths.shape,'  ', JHUrecovered.shape)

#JHUAC = JHUActive.groupby(['Country']).sum()
JHUCC = JHUconfirmed.groupby(['Country']).sum()
JHUDC = JHUdeaths.groupby(['Country']).sum()
JHURC = JHUrecovered.groupby(['Country']).sum()

# plotting global confirmed
JHUCNF1gr = JHUconfirmed1.groupby('date')['date', 'Confirmed'].sum().reset_index()
fig = px.line(JHUCNF1gr, x="date", y="Confirmed", title="Worldwide Confirmed over time", log_y=True)
fig.update_xaxes(showline=True, linewidth=2, linecolor='black', ticks="inside")
fig.update_yaxes(showline=True, linewidth=2, linecolor='black', ticks="inside", col=1)  #, tick0 = 0)
# fig.show()
fig.write_image('PLOT/code2/JHUConfirmedGlobal.png')

# plotting global Deceased
JHUDEC1gr = JHUdeaths1.groupby('date')['date', 'Deceased'].sum().reset_index()
fig = px.line(JHUDEC1gr, x="date", y="Deceased", title="Worldwide Deceased over time", log_y=True)
fig.update_xaxes(showline=True, linewidth=2, linecolor='black', ticks="inside")
fig.update_yaxes(showline=True, linewidth=2, linecolor='black', ticks="inside", col=1)
# fig.show()
fig.write_image('PLOT/code2/JHUDeceasedGlobal.png')

# plotting global Recovered
JHUREC1gr = JHUrecovered1.groupby('date')['date', 'Recovered'].sum().reset_index()
fig = px.line(JHUREC1gr, x="date", y="Recovered", title="Worldwide Recovered over time", log_y=True)
fig.update_xaxes(showline=True, linewidth=2, linecolor='black', ticks="inside")
fig.update_yaxes(showline=True, linewidth=2, linecolor='black', ticks="inside", col=1)
# fig.show()
fig.write_image('PLOT/code2/JHURecoveredGlobal.png')

# Load Indian Data from Indian Site
iRaw1 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/raw_data1.csv', index_col=None, parse_dates=True)
iRaw2 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/raw_data2.csv', index_col=None, parse_dates=True)
iRaw3 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/raw_data3.csv', index_col=None, parse_dates=True)
iRaw4 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/raw_data4.csv', index_col=None, parse_dates=True)
iRaw5 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/raw_data5.csv', index_col=None, parse_dates=True)
iState = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/state_wise.csv', index_col=None, parse_dates=True)
iStateD = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/state_wise_daily.csv', index_col=None, parse_dates=True)
iDist = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/district_wise.csv', index_col=None, parse_dates=True)
iCTS = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/case_time_series_1.csv', index_col=None, parse_dates=True)
idr1 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/death_and_recovered1.csv', index_col=None, parse_dates=True)
idr2 = pd.read_csv('DATASET/APIcovid19indiaorg/CSV/death_and_recovered2.csv', index_col=None, parse_dates=True)
idr = idr1.append(idr2)
iDeceased = idr.query('Patient_Status == "Deceased"')  # deceased
iRecovered = idr.query('Patient_Status == "Recovered"')  # recovered

iRaw11 = iRaw1.append(iRaw2)
iRaw31 = (iRaw3.append(iRaw4)).append(iRaw5)
iRaw11 = iRaw11.drop(["Estimated Onset Date","Backup Notes"], axis=1)
iRaw11 = iRaw11.rename(columns={"Patient Number":"ID", "Age Bracket":"Age","State Patient Number":"STNUM"})
iRaw11 = iRaw11.rename(columns={"Date Announced":"Date"})
iRaw11 = iRaw11.rename(columns={"Detected City":"City", "Detected District":"District", "Detected State": "State"})
iRaw31 = iRaw31.rename(columns={"Entry_ID":"ID", "Age Bracket":"Age","State Patient Number":"STNUM"})
iRaw31 = iRaw31.rename(columns={"Date Announced":"Date"})
iRaw31 = iRaw31.rename(columns={"Detected City":"City", "Detected District":"District", "Detected State": "State"})
iRaw31=iRaw31[iRaw31.columns[[0,1,2,3,4,5,6,7,8,10,12,11,16,17,18,13,14,15,9]]]
iRaw = iRaw11.append(iRaw31)

iCTS['Date'] = pd.to_datetime(iCTS['Date'])
iStateD['Date'] = pd.to_datetime(iStateD['Date']) 

# Trend for confirmed cases globally (JHU data)
threshold = 100
JHUCC1 = JHUCC.drop(["Lat", "Long"], axis=1).sort_values(JHUconfirmed.columns[-1], ascending=False)
plotGlobalGig(JHUCC1, threshold, "Confirmed", False, 0)
plt.savefig(out+'PLOT/code2/Global_Trend_Confirmed.png')
plt.show()
plt.close()

# Trend for Recovered cases globally (JHU data)
JHURC1 = JHURC.drop(["Lat", "Long"], axis=1).sort_values(JHUrecovered.columns[-1], ascending=False)
totConf = plotGlobalGig(JHUCC1, threshold, "Confirmed", True, 0)  # get total confirmed for % calc
plotGlobalGig(JHURC1, threshold, "Recovered", False, totConf)
plt.savefig(out+'PLOT/code2/Global_Trend_Recovered.png')
plt.show()
plt.close()

# Trend for Deceased cases globally (JHU data)
JHUDC1 = JHUDC.drop(["Lat", "Long"], axis=1).sort_values(JHUdeaths.columns[-1], ascending=False)
totConf = plotGlobalGig(JHUCC1, threshold, "Confirmed", True, 0)  # get total confirmed for % calc
plotGlobalGig(JHUDC1, threshold, "Deceased", False, totConf)
plt.savefig(out+'PLOT/code2/Global_Trend_Deceased.png')
plt.show()
plt.close()

###########################################################################################
###########################################################################################
###########################################################################################
#----------------------- Data for Kaggle competition --------------------------------------
kTrainData = pd.read_csv('Kaggle_Global_Forecast/train.csv', parse_dates = ['Date'])
print(kTrainData.head())
print(kTrainData.info())
kTestData = pd.read_csv('Kaggle_Global_Forecast/test.csv', parse_dates = ['Date'])
print(kTestData.head())
print(kTestData.info())
